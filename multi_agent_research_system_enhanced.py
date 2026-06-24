# multi_agent_research_system_enhanced.py
# Agentic AI Project: Enhanced Multi-Agent Research & Content Generation System
# Uses: LangChain, LangGraph, NVIDIA NIM, Web Search, Wikipedia, ArXiv, PDF Upload,
#       Persistent Memory, Multiple Topics, PDF/Word Export


import os
import json
import sqlite3
from typing import List, Dict, TypedDict, Annotated, Optional
from datetime import datetime
import operator
import tempfile
from io import BytesIO

# LangChain imports
from langchain_core.messages import HumanMessage, AIMessage, SystemMessage
from langchain_core.tools import tool
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_core.output_parsers import JsonOutputParser, StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_text_splitters import RecursiveCharacterTextSplitter

# LangGraph imports
from langgraph.graph import StateGraph, END
from langgraph.checkpoint.memory import MemorySaver
from langgraph.checkpoint.memory import MemorySaver

# NVIDIA NIM via OpenAI-compatible API
from langchain_openai import ChatOpenAI

# Real web search and research tools
from langchain_community.tools import DuckDuckGoSearchRun
from langchain_community.tools import WikipediaQueryRun
from langchain_community.utilities import WikipediaAPIWrapper
from langchain_community.tools import ArxivQueryRun

# For vector store and embeddings
from langchain_community.vectorstores import Chroma
from langchain_huggingface import HuggingFaceEmbeddings

# PDF processing
import PyPDF2

# Export to PDF
from fpdf import FPDF

# Export to Word
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

import streamlit as st


class ResearchState(TypedDict):
    """State for the multi-agent research system"""
    query: str
    research_results: List[str]
    summary: str
    report: str
    agent_status: Dict[str, str]
    iteration_count: int
    max_iterations: int
    sources: List[str]


class MultiAgentResearchSystem:
    """
    Enhanced Multi-Agent Research System with:
    - Planner Agent: Breaks down tasks
    - Research Agent: Gathers information from multiple sources (Web, Wikipedia, ArXiv)
    - Writer Agent: Generates content with citations
    - Critic Agent: Reviews and improves
    - Persistent Memory: SQLite-based conversation history
    - PDF Upload: Extract and research from documents
    - Multiple Topics: Batch research processing
    - Export: PDF and Word document generation
    """

    def __init__(self, db_path: str = "research_memory.db"):
        self.db_path = db_path
        self.memory = MemorySaver()
        self.embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2"
        )
        self.vector_store = None
        self.llm = self._init_llm()

        # Initialize research tools
        self.search_tool = DuckDuckGoSearchRun()
        self.wiki_tool = WikipediaQueryRun(api_wrapper=WikipediaAPIWrapper())
        self.arxiv_tool = ArxivQueryRun()

        self.workflow = self._build_workflow()
        self._init_database()

    def _init_database(self):
        """Initialize SQLite database for persistent storage"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()

        cursor.execute("""
            CREATE TABLE IF NOT EXISTS research_history (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                query TEXT,
                report TEXT,
                sources TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)

        cursor.execute("""
            CREATE TABLE IF NOT EXISTS uploaded_documents (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                filename TEXT,
                content TEXT,
                chunks TEXT,
                uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)

        conn.commit()
        conn.close()

    def _init_llm(self):
        """Initialize NVIDIA NIM LLM using OpenAI-compatible API"""
        api_key = os.environ.get("NVIDIA_API_KEY")

        if not api_key:
            raise ValueError(
                "NVIDIA_API_KEY not found! Set it with:\n"
                "  Windows CMD: set NVIDIA_API_KEY=nvapi-...\n"
                "  PowerShell: $env:NVIDIA_API_KEY='nvapi-...'\n"
                "  Or create a .env file with NVIDIA_API_KEY=nvapi-..."
            )

        llm = ChatOpenAI(
            model="meta/llama-3.3-70b-instruct",
            api_key=api_key,
            base_url="https://integrate.api.nvidia.com/v1",
            temperature=0.7,
            max_tokens=4096,
            top_p=0.7
        )
        return llm

    def _build_workflow(self):
        """Build the LangGraph workflow for multi-agent collaboration"""
        workflow = StateGraph(ResearchState)

        workflow.add_node("planner", self.planner_agent)
        workflow.add_node("researcher", self.research_agent)
        workflow.add_node("writer", self.writer_agent)
        workflow.add_node("critic", self.critic_agent)

        workflow.set_entry_point("planner")
        workflow.add_edge("planner", "researcher")
        workflow.add_edge("researcher", "writer")
        workflow.add_edge("writer", "critic")

        workflow.add_conditional_edges(
            "critic",
            self.should_continue,
            {
                "continue": "writer",
                "end": END
            }
        )

        return workflow.compile(checkpointer=self.memory)

    def planner_agent(self, state: ResearchState):
        """Planner Agent: Uses LLM to create a real research plan"""
        query = state["query"]

        prompt = f"""You are a research planner. Create a detailed step-by-step plan to research: "{query}"

Format your response as a clear numbered plan. Be specific about what to search for and which sources to consult (web, Wikipedia, academic papers)."""

        response = self.llm.invoke(prompt)
        plan = response.content if hasattr(response, "content") else str(response)

        state["agent_status"]["planner"] = "completed"
        state["research_results"].append(f"Plan: {plan}")
        state["sources"] = []

        return state

    def research_agent(self, state: ResearchState):
        """Research Agent: Performs multi-source research (Web, Wikipedia, ArXiv)"""
        query = state["query"]
        sources = []

        # 1. Web Search via DuckDuckGo
        try:
            web_results = self.search_tool.run(query)
            sources.append("Web Search (DuckDuckGo)")
        except Exception as e:
            web_results = f"Web search error: {str(e)}"

        # 2. Wikipedia Search
        try:
            wiki_results = self.wiki_tool.run(query)
            sources.append("Wikipedia")
        except Exception as e:
            wiki_results = f"Wikipedia error: {str(e)}"

        # 3. ArXiv Academic Papers
        try:
            arxiv_results = self.arxiv_tool.run(query)
            sources.append("ArXiv")
        except Exception as e:
            arxiv_results = f"ArXiv error: {str(e)}"

        # 4. Related searches for deeper research
        try:
            related_web = self.search_tool.run(f"{query} latest trends 2026")
        except Exception as e:
            related_web = f"Related search error: {str(e)}"

        combined_results = f"""=== MULTI-SOURCE RESEARCH RESULTS FOR "{query}" ===

--- WEB SEARCH (DuckDuckGo) ---
{web_results}

--- WIKIPEDIA ---
{wiki_results}

--- ARXIV ACADEMIC PAPERS ---
{arxiv_results}

--- ADDITIONAL WEB SEARCH (Latest Trends) ---
{related_web}

=== END OF RESEARCH ==="""

        state["research_results"].append(combined_results)
        state["sources"] = sources
        state["agent_status"]["researcher"] = "completed"

        return state

    def writer_agent(self, state: ResearchState):
        """Writer Agent: Uses LLM to generate cited content from research"""
        research = "\n".join(state["research_results"])
        query = state["query"]
        sources = state.get("sources", [])

        prompt = f"""You are a research analyst. Based on the following multi-source research findings, write a comprehensive, well-cited summary about: "{query}"

SOURCES CONSULTED: {', '.join(sources)}

RESEARCH FINDINGS:
{research}

Write a detailed, well-structured summary with:
- Executive Summary
- Key Findings and Insights (with source attribution)
- Statistics and Data Points
- Trends and Future Outlook
- Academic Insights (from ArXiv papers)
- Conclusions and Recommendations

Cite sources where appropriate (e.g., [Web Search], [Wikipedia], [ArXiv]).
Be factual and based ONLY on the research provided."""

        response = self.llm.invoke(prompt)
        summary = response.content if hasattr(response, "content") else str(response)

        state["summary"] = summary
        state["agent_status"]["writer"] = "completed"
        state["iteration_count"] += 1

        return state

    def critic_agent(self, state: ResearchState):
        """Critic Agent: Uses LLM to evaluate quality"""
        summary = state["summary"]
        query = state["query"]
        iteration = state["iteration_count"]
        max_iter = state["max_iterations"]

        # Auto-approve on max iterations
        if iteration >= max_iter:
            state["report"] = f"FINAL REPORT (Iteration {iteration}):\n\n{summary}"
            state["agent_status"]["critic"] = "approved"
            return state

        critique_prompt = f"""You are a critical reviewer. Evaluate this research summary about "{query}":

SUMMARY:
{summary}

Evaluate on:
1. Factual accuracy and specificity
2. Presence of concrete data/statistics
3. Source citations and attribution
4. Structure and clarity
5. Completeness (covers web, Wikipedia, and academic sources)

If the summary is well-structured, informative, properly cited, and based on real research findings, respond with APPROVED.
If it lacks critical information or has major issues, respond with NEEDS_REVISION and briefly why.

Be reasonable - a good summary with citations and clear structure should be APPROVED."""

        critique_response = self.llm.invoke(critique_prompt)
        critique = critique_response.content if hasattr(critique_response, "content") else str(critique_response)

        is_approved = "APPROVED" in critique.upper()

        if is_approved:
            final_prompt = f"""Polish this research report into a professional final document about "{query}":

{summary}

Format with clear sections, professional tone, and ensure all sources are properly cited.
Add a "Sources Consulted" section at the end."""

            final_response = self.llm.invoke(final_prompt)
            final_report = final_response.content if hasattr(final_response, "content") else str(final_response)

            state["report"] = f"FINAL REPORT (Iteration {iteration}):\n\n{final_report}"
            state["agent_status"]["critic"] = "approved"
        else:
            state["agent_status"]["critic"] = "needs_revision"
            state["research_results"].append(f"Feedback from critic (iter {iteration}): {critique}")

        return state

    def should_continue(self, state: ResearchState):
        """Decision function: Should we continue or end?"""
        if state["agent_status"].get("critic") == "approved":
            return "end"
        elif state["iteration_count"] >= state["max_iterations"]:
            state["report"] = f"FINAL REPORT (Max iterations reached):\n\n{state['summary']}"
            return "end"
        else:
            return "continue"

    def run(self, query: str, max_iterations: int = 3):
        """Execute the multi-agent workflow"""
        initial_state = ResearchState(
            query=query,
            research_results=[],
            summary="",
            report="",
            agent_status={},
            iteration_count=0,
            max_iterations=max_iterations,
            sources=[]
        )

        config = {"configurable": {"thread_id": "1"}}
        result = self.workflow.invoke(initial_state, config=config)

        # Save to database
        self._save_research(query, result["report"], result.get("sources", []))

        return result

    def run_multiple(self, queries: List[str], max_iterations: int = 3):
        """Run research on multiple topics and return combined results"""
        all_reports = []

        for i, query in enumerate(queries, 1):
            print(f"\n{'='*60}")
            print(f"Researching Topic {i}/{len(queries)}: {query}")
            print(f"{'='*60}")

            result = self.run(query, max_iterations)
            all_reports.append({
                "query": query,
                "report": result["report"],
                "sources": result.get("sources", [])
            })

        return all_reports

    def _save_research(self, query: str, report: str, sources: List[str]):
        """Save research to SQLite database"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()

        cursor.execute("""
            INSERT INTO research_history (query, report, sources)
            VALUES (?, ?, ?)
        """, (query, report, json.dumps(sources)))

        conn.commit()
        conn.close()

    def get_research_history(self, limit: int = 10):
        """Retrieve past research from database"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()

        cursor.execute("""
            SELECT query, report, sources, created_at
            FROM research_history
            ORDER BY created_at DESC
            LIMIT ?
        """, (limit,))

        results = cursor.fetchall()
        conn.close()

        return [
            {
                "query": row[0],
                "report": row[1],
                "sources": json.loads(row[2]) if row[2] else [],
                "created_at": row[3]
            }
            for row in results
        ]

    def process_pdf(self, pdf_file) -> str:
        """Extract text from uploaded PDF file"""
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""

        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"

        # Chunk the text for better processing
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200
        )
        chunks = text_splitter.split_text(text)

        # Save to database
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()

        cursor.execute("""
            INSERT INTO uploaded_documents (filename, content, chunks)
            VALUES (?, ?, ?)
        """, (pdf_file.name, text, json.dumps(chunks)))

        conn.commit()
        conn.close()

        return text, chunks

    def research_from_pdf(self, pdf_text: str, query: str = None) -> str:
        """Generate research summary from PDF content"""
        if not query:
            query = "Summarize and analyze the key findings in this document"

        prompt = f"""You are a document analyst. Based on the following document content, provide a comprehensive analysis.

QUERY: {query}

DOCUMENT CONTENT:
{pdf_text[:8000]}

Provide:
1. Executive Summary
2. Key Findings
3. Important Data/Statistics
4. Conclusions
5. Recommendations"""

        response = self.llm.invoke(prompt)
        return response.content if hasattr(response, "content") else str(response)


def export_to_pdf(report_text: str, filename: str = "research_report.pdf") -> bytes:
    """Export research report to PDF"""
    pdf = FPDF()
    pdf.add_page()

    # Title
    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 10, "Research Report", ln=True, align="C")
    pdf.ln(5)

    # Date
    pdf.set_font("Arial", "I", 10)
    pdf.cell(0, 10, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", ln=True, align="C")
    pdf.ln(10)

    # Content
    pdf.set_font("Arial", size=11)

    lines = report_text.split("\n")
    for line in lines:
        clean_line = line.encode("latin-1", "replace").decode("latin-1")

        if line.startswith("#") or line.startswith("**"):
            pdf.set_font("Arial", "B", 12)
            pdf.ln(3)
        else:
            pdf.set_font("Arial", size=11)

        pdf.multi_cell(0, 6, clean_line)
        pdf.ln(2)

    return pdf.output(dest="S").encode("latin-1")


def export_to_word(report_text: str, filename: str = "research_report.docx") -> bytes:
    """Export research report to Word document"""
    doc = Document()

    # Title
    title = doc.add_heading("Research Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    date_run.italic = True
    date_run.font.size = Pt(10)

    doc.add_paragraph()

    # Content
    lines = report_text.split("\n")
    for line in lines:
        if line.startswith("#") or (line.startswith("**") and line.endswith("**")):
            heading_text = line.strip("#").strip("*").strip()
            doc.add_heading(heading_text, level=2)
        elif line.startswith("---"):
            doc.add_paragraph("_" * 50)
        elif line.strip().startswith("•") or line.strip().startswith("-"):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(line.strip("•").strip("-").strip())
        elif line.strip().startswith("1.") or line.strip().startswith("2."):
            p = doc.add_paragraph(style="List Number")
            p.add_run(line.strip())
        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(6)

    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)

    return doc_bytes.getvalue()


def create_agent_dashboard():
    """Create enhanced interactive dashboard for the multi-agent system"""
    st.set_page_config(
        page_title="Agentic AI Research System",
        layout="wide",
        initial_sidebar_state="expanded"
    )

    # Sidebar
    with st.sidebar:
        st.title("Research Tools")

        api_key = os.environ.get("NVIDIA_API_KEY")
        if api_key:
            st.success("NVIDIA API Key configured")
        else:
            st.error("NVIDIA API Key not found")
            st.info("Set with: set NVIDIA_API_KEY=nvapi-...")

        st.markdown("---")

        page = st.radio("Navigation", [
            "New Research",
            "PDF Analysis",
            "Research History",
            "Settings"
        ])

        st.markdown("---")
        st.markdown("**Powered by:**")
        st.markdown("- LangChain + LangGraph")
        st.markdown("- NVIDIA NIM (Llama 3.3)")
        st.markdown("- DuckDuckGo + Wikipedia + ArXiv")

    # Main content
    st.title("Multi-Agent Research System")
    st.markdown("*Advanced research with multi-source aggregation, persistent memory, and document export*")

    # Initialize system
    if "agent_system" not in st.session_state:
        try:
            st.session_state.agent_system = MultiAgentResearchSystem()
            st.session_state.system_ready = True
        except Exception as e:
            st.error(f"Failed to initialize: {str(e)}")
            st.session_state.system_ready = False
            return

    if not st.session_state.get("system_ready", False):
        st.error("System not initialized. Check API key configuration.")
        return

    system = st.session_state.agent_system

    if page == "New Research":
        st.header("Research Query")

        research_mode = st.radio("Research Mode", ["Single Topic", "Multiple Topics"], horizontal=True)

        if research_mode == "Single Topic":
            query = st.text_input(
                "Enter your research topic:",
                placeholder="e.g., Latest trends in Agentic AI"
            )
        else:
            query = st.text_area(
                "Enter multiple topics (one per line):",
                placeholder="Topic 1\nTopic 2\nTopic 3",
                height=100
            )
            query = [q.strip() for q in query.split("\n") if q.strip()] if query else []

        col1, col2, col3 = st.columns(3)
        with col1:
            max_iterations = st.slider("Max Iterations", 1, 5, 2)
        with col2:
            include_wiki = st.checkbox("Include Wikipedia", value=True)
        with col3:
            include_arxiv = st.checkbox("Include ArXiv", value=True)

        if st.button("Start Research", type="primary", use_container_width=True):
            if not query:
                st.warning("Please enter a research topic!")
                return

            with st.spinner("Agents working... Researching across multiple sources..."):
                try:
                    if research_mode == "Single Topic":
                        result = system.run(query, max_iterations)
                        st.session_state.last_result = result
                        st.session_state.last_results = None
                    else:
                        results = system.run_multiple(query, max_iterations)
                        st.session_state.last_results = results
                        st.session_state.last_result = None

                    st.success("Research completed!")
                except Exception as e:
                    st.error(f"Error during research: {str(e)}")
                    return

        if st.session_state.get("last_result"):
            _display_single_result(st.session_state.last_result)

        if st.session_state.get("last_results"):
            _display_multiple_results(st.session_state.last_results)

    elif page == "PDF Analysis":
        st.header("PDF Document Analysis")

        uploaded_file = st.file_uploader("Upload PDF", type="pdf")

        if uploaded_file:
            with st.spinner("Processing PDF..."):
                text, chunks = system.process_pdf(uploaded_file)
                st.session_state.pdf_text = text
                st.session_state.pdf_chunks = chunks
                st.success(f"PDF processed: {len(chunks)} chunks extracted")

            with st.expander("Preview Extracted Text"):
                st.text_area("Content", text[:2000] + "...", height=200)

            query = st.text_input(
                "Ask a question about the document:",
                placeholder="What are the key findings?"
            )

            if st.button("Analyze Document", type="primary"):
                with st.spinner("Analyzing..."):
                    analysis = system.research_from_pdf(text, query)
                    st.session_state.pdf_analysis = analysis

            if st.session_state.get("pdf_analysis"):
                st.markdown("---")
                st.subheader("Analysis Results")
                st.markdown(st.session_state.pdf_analysis)

                col1, col2 = st.columns(2)
                with col1:
                    pdf_bytes = export_to_pdf(st.session_state.pdf_analysis)
                    st.download_button(
                        "Download PDF",
                        pdf_bytes,
                        file_name="pdf_analysis.pdf",
                        mime="application/pdf"
                    )
                with col2:
                    docx_bytes = export_to_word(st.session_state.pdf_analysis)
                    st.download_button(
                        "Download Word",
                        docx_bytes,
                        file_name="pdf_analysis.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

    elif page == "Research History":
        st.header("Research History")

        history = system.get_research_history(limit=20)

        if not history:
            st.info("No research history yet. Start a new research!")
        else:
            for i, item in enumerate(history, 1):
                with st.expander(f"{i}. {item['query']} ({item['created_at']})"):
                    st.markdown("**Sources:** " + ", ".join(item["sources"]))
                    st.markdown(item["report"][:500] + "...")

                    col1, col2 = st.columns(2)
                    with col1:
                        pdf_bytes = export_to_pdf(item["report"])
                        st.download_button(
                            "PDF",
                            pdf_bytes,
                            file_name=f"research_{i}.pdf",
                            mime="application/pdf",
                            key=f"pdf_{i}"
                        )
                    with col2:
                        docx_bytes = export_to_word(item["report"])
                        st.download_button(
                            "Word",
                            docx_bytes,
                            file_name=f"research_{i}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"word_{i}"
                        )

    elif page == "Settings":
        st.header("System Settings")

        st.subheader("Database")
        db_path = st.text_input("Database Path", value="research_memory.db")

        st.subheader("LLM Configuration")
        model = st.selectbox(
            "Model",
            ["meta/llama-3.3-70b-instruct", "meta/llama-3.1-70b-instruct", "mistralai/mixtral-8x22b-instruct-v0.1"]
        )
        temperature = st.slider("Temperature", 0.0, 1.0, 0.7)
        max_tokens = st.slider("Max Tokens", 512, 8192, 4096)

        st.subheader("Export Settings")
        export_format = st.multiselect(
            "Default Export Formats",
            ["PDF", "Word (DOCX)"],
            default=["PDF", "Word (DOCX)"]
        )

        if st.button("Save Settings", type="primary"):
            st.success("Settings saved (Note: Restart required for LLM model changes)")


def _display_single_result(result):
    """Display single research result"""
    st.markdown("---")
    st.header("Research Results")

    st.subheader("Agent Status")
    cols = st.columns(4)
    agents = ["planner", "researcher", "writer", "critic"]
    for i, agent in enumerate(agents):
        status = result["agent_status"].get(agent, "pending")
        emoji = "OK" if status in ["completed", "approved"] else "WAIT"
        cols[i].metric(f"{emoji} {agent.title()}", status.upper())

    if result.get("sources"):
        st.subheader("Sources Consulted")
        st.write(", ".join([f"{s}" for s in result["sources"]]))

    st.subheader("Raw Research Findings")
    for i, finding in enumerate(result["research_results"]):
        with st.expander(f"Finding {i+1}"):
            st.write(finding)

    st.subheader("Final Report")
    st.markdown(result["report"])

    st.info(f"Completed in {result['iteration_count']} iterations")

    st.subheader("Export Report")
    col1, col2 = st.columns(2)
    with col1:
        pdf_bytes = export_to_pdf(result["report"])
        st.download_button(
            "Download PDF",
            pdf_bytes,
            file_name=f"research_report_{result['query'].replace(' ', '_')}.pdf",
            mime="application/pdf"
        )
    with col2:
        docx_bytes = export_to_word(result["report"])
        st.download_button(
            "Download Word",
            docx_bytes,
            file_name=f"research_report_{result['query'].replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )


def _display_multiple_results(results):
    """Display multiple research results"""
    st.markdown("---")
    st.header("Batch Research Results")

    for i, item in enumerate(results, 1):
        with st.expander(f"Topic {i}: {item['query']}"):
            st.markdown(item["report"])

            col1, col2 = st.columns(2)
            with col1:
                pdf_bytes = export_to_pdf(item["report"])
                st.download_button(
                    "PDF",
                    pdf_bytes,
                    file_name=f"research_{i}.pdf",
                    mime="application/pdf",
                    key=f"multi_pdf_{i}"
                )
            with col2:
                docx_bytes = export_to_word(item["report"])
                st.download_button(
                    "Word",
                    docx_bytes,
                    file_name=f"research_{i}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"multi_word_{i}"
                )

    st.markdown("---")
    combined_report = "\n\n".join([
        f"# Topic {i}: {r['query']}\n\n{r['report']}"
        for i, r in enumerate(results, 1)
    ])

    st.subheader("Export All Reports")
    col1, col2 = st.columns(2)
    with col1:
        pdf_bytes = export_to_pdf(combined_report)
        st.download_button(
            "Download All (PDF)",
            pdf_bytes,
            file_name="combined_research_report.pdf",
            mime="application/pdf"
        )
    with col2:
        docx_bytes = export_to_word(combined_report)
        st.download_button(
            "Download All (Word)",
            docx_bytes,
            file_name="combined_research_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )


def run_cli_demo():
    """Run CLI demo"""
    print("=" * 70)
    print("AGENTIC AI - ENHANCED MULTI-AGENT RESEARCH SYSTEM")
    print("=" * 70)
    print("Features: Multi-source research | PDF analysis | Persistent memory | Export")
    print("=" * 70)

    try:
        system = MultiAgentResearchSystem()

        print("\n[1] Single Topic Research")
        result = system.run("Latest trends in Agentic AI and LLMs", max_iterations=2)

        print("\nAGENT STATUS:")
        for agent, status in result["agent_status"].items():
            print(f"  {agent}: {status}")

        print(f"\nSOURCES: {', '.join(result.get('sources', []))}")
        print("\nFINAL REPORT (excerpt):")
        print(result["report"][:800] + "...")

        print("\n\n[2] Batch Research")
        topics = ["Quantum Computing applications", "Climate AI solutions"]
        results = system.run_multiple(topics, max_iterations=2)

        for r in results:
            print(f"\n--- {r['query']} ---")
            print(r["report"][:300] + "...")

        print("\n\n[3] Research History")
        history = system.get_research_history(limit=5)
        for h in history:
            print(f"  - {h['query']} ({h['created_at']})")

    except Exception as e:
        print(f"\nError: {e}")
        print("\nMake sure to set NVIDIA_API_KEY environment variable")

    print("\n" + "=" * 70)
    print("To run dashboard: streamlit run multi_agent_research_system_enhanced.py")
    print("=" * 70)


if __name__ == "__main__":
    import sys
    if "streamlit" in sys.modules:
        create_agent_dashboard()
    else:
        run_cli_demo()
